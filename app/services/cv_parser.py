"""
CV/Resume Parser Service
Extracts text from PDF and DOCX files for RAG-based interview personalization.
"""

import os
import re
from PyPDF2 import PdfReader
from docx import Document


def extract_text_from_pdf(file_path):
    """Extract text content from a PDF file."""
    try:
        reader = PdfReader(file_path)
        text_parts = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
        return "\n".join(text_parts)
    except Exception as e:
        print(f"Error extracting PDF text: {e}")
        raise ValueError(f"Failed to parse PDF: {str(e)}")


def extract_text_from_docx(file_path):
    """Extract text content from a DOCX file."""
    try:
        doc = Document(file_path)
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)
        return "\n".join(text_parts)
    except Exception as e:
        print(f"Error extracting DOCX text: {e}")
        raise ValueError(f"Failed to parse DOCX: {str(e)}")


def clean_cv_text(text):
    """Clean extracted CV text — remove excessive whitespace, normalize formatting."""
    # Remove excessive blank lines
    text = re.sub(r'\n{3,}', '\n\n', text)
    # Remove excessive spaces
    text = re.sub(r'  +', ' ', text)
    # Strip each line
    lines = [line.strip() for line in text.split('\n')]
    text = '\n'.join(lines)
    return text.strip()


def parse_cv(file_path):
    """
    Parse a CV file (PDF or DOCX) and return cleaned text.
    
    Args:
        file_path: Path to the CV file
    
    Returns:
        str: Extracted and cleaned CV text
    
    Raises:
        ValueError: If file format is unsupported or parsing fails
    """
    if not os.path.exists(file_path):
        raise ValueError(f"File not found: {file_path}")

    ext = os.path.splitext(file_path)[1].lower()
    
    if ext == '.pdf':
        raw_text = extract_text_from_pdf(file_path)
    elif ext == '.docx':
        raw_text = extract_text_from_docx(file_path)
    else:
        raise ValueError(f"Unsupported file format: {ext}. Only PDF and DOCX are supported.")

    if not raw_text or len(raw_text.strip()) < 50:
        raise ValueError("CV appears to be empty or contains too little text. Please upload a valid CV.")

    return clean_cv_text(raw_text)


def build_cv_context(cv_text):
    """
    Build a context block from CV text to inject into the interview system prompt.
    
    Args:
        cv_text: Extracted CV text
    
    Returns:
        str: Formatted context block for system prompt injection
    """
    return f"""

--- CANDIDATE'S CV/RESUME (Uploaded by the candidate) ---
{cv_text}
--- END OF CV ---

IMPORTANT INSTRUCTIONS FOR USING THE CV:
1. Ask personalized questions about the candidate's specific projects, skills, and experience mentioned in the CV.
2. Probe deeper into claims made in the CV — ask for details, challenges faced, and outcomes.
3. Verify consistency between CV claims and verbal responses during the interview.
4. Tailor technical questions to the technologies, tools, and skills listed in the CV.
5. Ask about gaps or inconsistencies you notice in the CV.
6. Do NOT repeat information from the CV back to the candidate — instead, ask them to elaborate on specific items.
7. Use the CV as a guide to structure your questions, but still cover all required interview areas.
"""


def allowed_file(filename, allowed_extensions):
    """Check if a filename has an allowed extension."""
    if '.' not in filename:
        return False
    ext = filename.rsplit('.', 1)[1].lower()
    clean_allowed = {e.lstrip('.').lower() for e in allowed_extensions}
    return ext in clean_allowed
